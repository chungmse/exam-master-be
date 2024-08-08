import os, io
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from docx import Document
from typing import List, Dict, Tuple, Union
from pydantic import BaseModel
from db import db
import uuid
import time
router = APIRouter(prefix="/question", tags=["question"])


class QuestionData(BaseModel):
    subject_id: int
    question_text: str
    answer: int
    option1: str
    option2: str
    option3: str
    option4: str
    mark: float
    unit: str
    mix: int


Images_dir = "public/img"
if not os.path.exists(Images_dir):
    os.makedirs(Images_dir)


class ImportRequest(BaseModel):
    file_name: str


class ErrorDetail(BaseModel):
    table_index: int
    row_index: int
    message: str


@router.post("/")
def question_root(db_conn: db.get_db = Depends()):
    cursor = db_conn.cursor()
    cursor.execute("SELECT 'Another example route!' as message")
    result = cursor.fetchone()
    return {"message": result["message"]}


def validate_data(question_data) -> Tuple[bool, str]:
    # Kiểm tra kiểu dữ liệu của từng ô
    if question_data[0] is None or not isinstance(question_data[0], int):  # subject_id
        return False, "Câu hỏi phải đặt thành số"
    if (
        question_data[1] is None
        or not isinstance(question_data[1], str)
        or not question_data[1]
    ):  # question_text
        return False, "Nội dung câu hỏi không được để trống"
    if (
        question_data[2] is None
        or not isinstance(question_data[2], int)
        or question_data[2] not in [1, 2, 3, 4]
    ):  # answer
        return False, "Câu hỏi phải đúng kí tự (A, B, C, or D)"
    for i in range(3, 7):  # option1 to option4
        if (
            question_data[i] is None
            or not isinstance(question_data[i], str)
            or not question_data[i]
        ):
            return False, f"Đáp án {i - 2} không được để trống"
    if question_data[7] is None or not isinstance(question_data[7], float):  # mark
        return False, "Điểm phải là giá trị số thập phân"
    if (
        question_data[8] is None
        or not isinstance(question_data[8], str)
        or not question_data[8]
    ):  # unit
        return False, "Chương nội dung không đươc để trống"
    if (
        question_data[9] is None
        or not isinstance(question_data[9], int)
        or question_data[9] not in [0, 1]
    ):  # mix
        return False, "Chỉ lựa chọn có hoặc không"

    return True, ""


subject = None
lecturer = None
date = None
expected_number_of_quiz = None


def process_file(file_contents) -> Tuple[List[QuestionData], List[Dict]]:
    doc = Document(io.BytesIO(file_contents))
    data = []
    errors = []
    # subject_id = 1
    question_count = 0

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Subject:"):
            global subject
            subject = paragraph.text.split(":")[1].strip()
        elif paragraph.text.startswith("Number of Quiz:"):
            global expected_number_of_quiz
            expected_number_of_quiz = int(paragraph.text.split(":")[1].strip())
        elif paragraph.text.startswith("Lecturer:"):
            global lecturer
            lecturer = paragraph.text.split(":")[1].strip()
        elif paragraph.text.startswith("Date:"):
            global date
            date = paragraph.text.split(":")[1].strip()

    for table_index, table in enumerate(doc.tables):
        question_data = [None] * 10  # Tạo danh sách để chứa thông tin câu hỏi
        # question_data[0] = subject_id
        question_count += 1
        for row_index, row in enumerate(table.rows):
            cell_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cell_data.append(cell_text)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.element.drawing_lst:
                            for drawing in run.element.drawing_lst:
                                for blip in drawing.xpath(".//a:blip"):
                                    rId = blip.get(
                                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                    )
                                    image_part = doc.part.related_parts[rId]
                                    image_bytes = image_part.blob
                                    image_name = (
                                        f"{uuid.uuid4().hex[:8]}_{int(time.time())}.png"
                                    )
                                    image_path = os.path.join(Images_dir, image_name)
                                    with open(image_path, "wb") as img_file:
                                        img_file.write(image_bytes)
                                    cell_data[-1] += f" [img:{image_name}]"
            if row_index == 0:
                question_data[0] = table_index + 1
                question_data[1] = cell_data[1]  # Nội dung câu hỏi
            elif row_index == 1:
                question_data[3] = cell_data[1]  # Đáp án a
            elif row_index == 2:
                question_data[4] = cell_data[1]  # Đáp án b
            elif row_index == 3:
                question_data[5] = cell_data[1]  # Đáp án c
            elif row_index == 4:
                question_data[6] = cell_data[1]  # Đáp án d
            elif row_index == 5:
                answer = cell_data[1].lower() if cell_data[1] else None
                if answer == "a":
                    question_data[2] = 1
                elif answer == "b":
                    question_data[2] = 2
                elif answer == "c":
                    question_data[2] = 3
                elif answer == "d":
                    question_data[2] = 4
                else:
                    question_data[2] = None  # Đáp án đúng
            elif row_index == 6:
                try:
                    question_data[7] = (
                        float(cell_data[1]) if cell_data[1] else None
                    )  # Số điểm cho câu trả lời đúng
                except ValueError:
                    question_data[7] = None
            elif row_index == 7:
                question_data[8] = cell_data[1]  # Chương nội dung
            elif row_index == 8:
                question_data[9] = (
                    1 if cell_data[1].lower() == "yes" else 0
                )  # Kiểu câu hỏi

        valid, message = validate_data(question_data)
        if valid:
            data.append(
                QuestionData(
                    subject_id=question_data[0],
                    question_text=question_data[1],
                    answer=question_data[2],
                    option1=question_data[3],
                    option2=question_data[4],
                    option3=question_data[5],
                    option4=question_data[6],
                    mark=question_data[7],
                    unit=question_data[8],
                    mix=question_data[9],
                )
            )

        else:
            errors.append(
                {
                    "table_index": table_index + 1,
                    "row_index": row_index + 1,
                    "message": message,
                }
            )
    if (
        expected_number_of_quiz is not None
        and question_count != expected_number_of_quiz
    ):
        errors.append(
            {
                "table_index": None,
                "row_index": None,
                "message": f"Cài đặt {expected_number_of_quiz} câu hỏi, nhưng chỉ tìm thấy {question_count} bảng dữ liệu",
            }
        )
    return data, errors  # Move this line outside the loop


@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        return {"err": True, "msg": "Chỉ chấp nhận file .docx"}

    contents = await file.read()
    processed_data, errors = process_file(contents)

    if not processed_data and not errors:
        return {"err": True, "msg": "Dữ liệu không đúng cấu trúc"}

    if errors:
        msgText = "Lỗi được tìm thấy:"
        for error in errors:
            msgText += " " + error["message"] + "."
        return {"err": True, "msg": msgText}

    # Check if subject is already existed
    cursor = db.get_db().cursor()
    cursor.execute(
        """
        SELECT * FROM subjects WHERE subject_name = ?
    """,
        (subject,),
    )
    subject_data = cursor.fetchone()
    if subject_data is None:
        return {"err": True, "msg": "Môn học không tồn tại"}

    return {
        "subject": subject,
        "number_of_questions": expected_number_of_quiz,
        "date": date,
        "list_questions": processed_data,
    }


@router.post("/import")
def question_import(processed_data: List[QuestionData], db_conn: db.get_db = Depends()):
    try:
        cursor = db_conn.cursor()

        for question_data in processed_data:
            sql = """
                    INSERT INTO questions (
                        subject_id, question_text, answer, option1, option2, option3, option4, mark, unit, mix
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
            cursor.execute(
                sql,
                (
                    question_data.subject_id,
                    question_data.question_text,
                    question_data.answer,
                    question_data.option1,
                    question_data.option2,
                    question_data.option3,
                    question_data.option4,
                    question_data.mark,
                    question_data.unit,
                    question_data.mix,
                ),
            )

        db_conn.commit()

        return {"message": "Thêm dữ liệu thành công"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))